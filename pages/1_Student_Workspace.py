# pages/1_Student_Workspace.py
import io, datetime, time, html as _html
import streamlit as st
from streamlit_quill import st_quill
from streamlit.components.v1 import html as st_html

from lib.ui import inject_css, md_to_html, html_to_text
from lib.clients import get_config, get_llm_client
from lib.storage import (
    get_or_create_worksheets, append_row_safe, save_draft_row, load_last_draft,
    log_turn_row
)

st.set_page_config(page_title="Student Workspace", layout="wide")
inject_css()

cfg = get_config()
EVENTS_WS, DRAFTS_WS = get_or_create_worksheets()

# Session defaults
st.session_state.setdefault("assignment_id", cfg["ASSIGNMENT_DEFAULT"])
st.session_state.setdefault("chat", [])
st.session_state.setdefault("llm_outputs", [])
st.session_state.setdefault("draft_html", "")
st.session_state.setdefault("report", None)
st.session_state.setdefault("last_saved_at", None)
st.session_state.setdefault("last_autosave_at", None)
st.session_state.setdefault("last_saved_html", "")
st.session_state.setdefault("pending_prompt", None)

# Require auth
if not st.session_state.get("__auth_ok"):
    st.error("Please login from Home to use the workspace.")
    st.stop()

LLM = get_llm_client()

def excerpt(text, n=300):
    t = text or ""
    return t if len(t) <= n else t[:n] + " …"

# --- Similarity helpers (SBERT → TF-IDF → difflib) ---
SIM_THRESHOLD = cfg["SIM_THRESHOLD"]
SIM_BACKEND = "none"
try:
    from sentence_transformers import SentenceTransformer, util as sbert_util
    @st.cache_resource
    def _load_sbert():
        return SentenceTransformer("all-MiniLM-L6-v2")
    _SBERT = _load_sbert()
    SIM_BACKEND = "sbert"
except Exception:
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        SIM_BACKEND = "tfidf"
    except Exception:
        from difflib import SequenceMatcher
        SIM_BACKEND = "difflib"

def _segment(text: str):
    return [p.strip() for p in (text or "").split("\n") if p.strip()]

def compute_similarity_report(final_text, llm_texts, sim_thresh=SIM_THRESHOLD):
    finals = _segment(final_text)
    llm_segs = [s for t in llm_texts for s in _segment(t)]
    if not finals or not llm_segs:
        return {"backend": SIM_BACKEND, "mean": 0.0, "high_share": 0.0, "rows": []}

    rows, high_tokens = [], 0
    total_tokens = sum(len(s.split()) for s in finals)

    if SIM_BACKEND == "sbert":
        Ef = _SBERT.encode(finals, convert_to_tensor=True, normalize_embeddings=True)
        El = _SBERT.encode(llm_segs, convert_to_tensor=True, normalize_embeddings=True)
        sims = sbert_util.cos_sim(Ef, El).cpu().numpy()
        for i, fseg in enumerate(finals):
            j = int(sims[i].argmax()); s = float(sims[i, j]); nearest = llm_segs[j]
            rows.append({"final_seg": excerpt(fseg, 200), "nearest_llm": excerpt(nearest, 200), "cosine": round(s, 3)})
            if s >= sim_thresh: high_tokens += len(fseg.split())
    elif SIM_BACKEND == "tfidf":
        vectorizer = TfidfVectorizer().fit(finals + llm_segs)
        F = vectorizer.transform(finals); L = vectorizer.transform(llm_segs)
        sims = cosine_similarity(F, L)
        for i, fseg in enumerate(finals):
            j = int(sims[i].argmax()); s = float(sims[i, j]); nearest = llm_segs[j]
            rows.append({"final_seg": excerpt(fseg, 200), "nearest_llm": excerpt(nearest, 200), "cosine": round(s, 3)})
            if s >= sim_thresh: high_tokens += len(fseg.split())
    else:
        def cos_like(a,b): return SequenceMatcher(None,a,b).ratio()
        for fseg in finals:
            best, near = 0.0, ""
            for l in llm_segs:
                c = cos_like(fseg,l)
                if c>best: best, near = c, l
            rows.append({"final_seg": excerpt(fseg,200), "nearest_llm": excerpt(near,200), "cosine": round(best,3)})
            if best >= sim_thresh: high_tokens += len(fseg.split())

    mean_sim = 0.0 if not rows else round(sum(r["cosine"] for r in rows) / len(rows), 3)
    high_share = round(high_tokens / max(1, total_tokens), 3)
    return {"backend": SIM_BACKEND, "mean": mean_sim, "high_share": high_share, "rows": rows[:30]}

# --- LLM ---
def ask_llm(prompt_text: str):
    chunks = []
    try:
        for ch in LLM.generate_content([prompt_text], stream=True):
            if getattr(ch, "text", None):
                chunks.append(ch.text)
    except Exception as e:
        chunks.append(f"Error: {e}")
    return "".join(chunks)

def maybe_autosave():
    now = time.time()
    last_ts = st.session_state.get("last_autosave_at") or 0
    changed = (st.session_state.draft_html or "") != (st.session_state.last_saved_html or "")
    if changed and (now - last_ts) >= cfg["AUTO_SAVE_SECONDS"]:
        save_draft_row(DRAFTS_WS,
                       st.session_state.user_id,
                       st.session_state.assignment_id,
                       st.session_state.draft_html)
        st.session_state["last_autosave_at"] = now

# --- Header ---
st.markdown(
    f"""
<div class="header-bar">
  <div class="status-chip">User: {st.session_state.get('user_id')}</div>
  <div class="status-chip">Assignment: {st.session_state.get('assignment_id')}</div>
  <div class="status-chip">Similarity backend: {SIM_BACKEND}</div>
  <div class="small-muted">Last saved: {st.session_state['last_saved_at'].strftime("%H:%M:%S") if st.session_state.get('last_saved_at') else "—"}</div>
</div>
""",
    unsafe_allow_html=True,
)

# --- Toolbar ---
t1, t2, t3 = st.columns([1.2, 0.9, 0.8])
with t1:
    st.session_state["assignment_id"] = st.text_input("Assignment ID", value=st.session_state["assignment_id"])
with t2:
    if st.button("🔄 Load Last Draft", use_container_width=True):
        html = load_last_draft(DRAFTS_WS, st.session_state["user_id"], st.session_state["assignment_id"])
        if html:
            st.session_state["draft_html"] = html
            st.success("Loaded last saved draft.")
            st.rerun()
        else:
            st.warning("No saved draft found.")
with t3:
    if st.button("🧹 Clear Chat", use_container_width=True):
        st.session_state["chat"] = []
        st.session_state["llm_outputs"] = []
        st.toast("Chat cleared")

left, right = st.columns([0.5, 0.5], gap="large")

# --- Left: Assistant ---
with left:
    st.subheader("💬 Assistant")

    if not st.session_state["chat"]:
        bubbles_html = '<div class="chat-empty">Ask for ideas, critique, or examples.</div>'
    else:
        out = []
        for m in st.session_state["chat"]:
            css = "chat-user" if m["role"] == "user" else "chat-assistant"
            content = md_to_html(m["text"]) if m["role"] == "assistant" \
                     else _html.escape(m["text"]).replace("\n", "<br>")
            out.append(f'<div class="chat-bubble {css}">{content}</div>')
        bubbles_html = "".join(out)

    if st.session_state.get("pending_prompt"):
        bubbles_html += '<div class="chat-bubble chat-assistant">…thinking</div>'

    st_html(
        f'<div id="chatbox" class="chat-box">{bubbles_html}</div>'
        f'<script>var b=document.getElementById("chatbox"); if(b) b.scrollTop=b.scrollHeight;</script>',
        height=600
    )

    with st.form("chat_form", clear_on_submit=True):
        c1, c2 = st.columns([4, 1])
        with c1:
            prompt = st.text_input("Ask…", "", placeholder="Type and press Send", label_visibility="collapsed")
        with c2:
            send = st.form_submit_button("Send")

    if send and (prompt or "").strip():
        st.session_state["chat"].append({"role": "user", "text": prompt})
        st.session_state["pending_prompt"] = prompt
        st.rerun()

    if st.session_state.get("pending_prompt"):
        with st.spinner("Generating response…"):
            p = st.session_state["pending_prompt"]
            st.session_state["pending_prompt"] = None
            reply = ask_llm(p)
            st.session_state["chat"].append({"role": "assistant", "text": reply})
            st.session_state["llm_outputs"].append(reply)

            # Log the single consolidated turn
            log_turn_row(EVENTS_WS,
                         st.session_state["user_id"],
                         st.session_state["assignment_id"],
                         p, reply,
                         turn=sum(1 for m in st.session_state["chat"] if m["role"] == "user"))
        st.rerun()

# --- Right: Draft ---
with right:
    st.subheader("📝 Draft")
    # Rich editor
    try:
        out = st_quill(value=st.session_state["draft_html"], placeholder="Write your draft here…", html=True, key="editor")
        if isinstance(out, dict) and out.get("html"):
            st.session_state["draft_html"] = out["html"]
        elif isinstance(out, str) and out:
            st.session_state["draft_html"] = out
    except TypeError:
        st.session_state["draft_html"] = st_quill(st.session_state["draft_html"])

    maybe_autosave()

    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("💾 Save Draft", use_container_width=True):
            save_draft_row(DRAFTS_WS, st.session_state["user_id"],
                           st.session_state["assignment_id"], st.session_state["draft_html"])
            st.session_state["last_saved_at"] = datetime.datetime.now()
            st.session_state["last_saved_html"] = st.session_state["draft_html"]
            st.toast("Draft saved")

    with c2:
        if st.button("📊 Run Similarity", use_container_width=True):
            plain = html_to_text(st.session_state["draft_html"])
            if plain.strip() and st.session_state["llm_outputs"]:
                report = compute_similarity_report(plain, st.session_state["llm_outputs"], SIM_THRESHOLD)
                st.session_state["report"] = report
                st.success(f"Mean: {report['mean']} | High-sim: {report['high_share']*100:.1f}%")
                with st.expander("Matches (trimmed)"):
                    for r in report["rows"]:
                        st.markdown(f"- **Cos:** {r['cosine']}  \n  **Final:** {r['final_seg']}  \n  **LLM:** {r['nearest_llm']}")
            else:
                st.warning("Need draft text + at least one LLM response.")

    with c3:
        if st.button("⬇️ Export Evidence (DOCX)", use_container_width=True):
            try:
                import docx
                def export_docx(user_id, assign_id, chat, draft_html, report):
                    final_text = html_to_text(draft_html)
                    d = docx.Document()
                    d.add_heading("Coursework Evidence Pack", 0)
                    d.add_paragraph(f"User ID: {user_id}")
                    d.add_paragraph(f"Assignment ID: {assign_id}")
                    d.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    d.add_heading("Chat with LLM", level=1)
                    for m in chat:
                        who = "Student" if m["role"] == "user" else "LLM"
                        d.add_paragraph(f"{who}: {m['text']}")
                    d.add_heading("Final Draft (plain text extract)", level=1)
                    for para in final_text.split("\n"):
                        d.add_paragraph(para)
                    rep = report or {"backend":"-","mean":0.0,"high_share":0.0,"rows":[]}
                    d.add_heading("Similarity Report", level=1)
                    d.add_paragraph(f"Backend: {rep.get('backend','-')}")
                    d.add_paragraph(f"Mean similarity: {rep.get('mean',0.0)}")
                    d.add_paragraph(f"High-sim share: {rep.get('high_share',0.0)*100:.1f}%")
                    for r in rep.get("rows", []):
                        d.add_paragraph(f"- Cosine: {r['cosine']} | Final: {r['final_seg']} | LLM: {r['nearest_llm']}")
                    buf = io.BytesIO(); d.save(buf); buf.seek(0); return buf.read()

                data = export_docx(st.session_state["user_id"], st.session_state["assignment_id"],
                                   st.session_state["chat"], st.session_state["draft_html"],
                                   st.session_state.get("report"))
                st.download_button(
                    "Download DOCX",
                    data=data,
                    file_name=f"evidence_{st.session_state['user_id']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Export failed: {e}")
